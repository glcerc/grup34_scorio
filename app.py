import os
import io
import re
import json
import docx
import PyPDF2
import time
import requests
import streamlit as st
import google.generativeai as genai

from dotenv import load_dotenv
from pymongo import MongoClient
from datetime import datetime, timedelta
from bson import ObjectId


# Gemini AI konfigürasyonunu da ekleyin (main() fonksiyonundan önce):
# Gemini AI yapılandırması
try:
    genai.configure(api_key=os.getenv("GEMINI_API_KEY"))
    model = genai.GenerativeModel('gemini-2.0-flash')
except Exception as e:
    st.error(f"❌ Gemini API yapılandırma hatası: {e}")


# Sayfa ayarları
st.set_page_config(
    page_title="Essay Grader AI",
    page_icon="📝",
    layout="wide"
)

# Çevre değişkenlerini yükle
load_dotenv()

# MongoDB bağlantısı
@st.cache_resource
def init_mongodb():
    uri = os.getenv('MONGO_URI')
    if not uri:
        st.error("⚠️ MONGO_URI çevresel değişkeni tanımlı değil!")
        raise Exception("MONGO_URI eksik")
    try:
        client = MongoClient(
            uri,
            serverSelectionTimeoutMS=10000,
            connectTimeoutMS=10000,
            socketTimeoutMS=10000,
            maxPoolSize=1
        )
        db = client['essay_grader']
        # Bağlantıyı test et
        
        collections = db.list_collection_names()
       
        return db
    except Exception as e:
        
        raise



def extract_text_from_file(uploaded_file):
    """Dosyadan metin çıkarma"""
    try:
        file_extension = uploaded_file.name.split('.')[-1].lower()
        
        if file_extension == 'txt':
            return uploaded_file.read().decode('utf-8')
        
        elif file_extension == 'pdf':
            pdf_reader = PyPDF2.PdfReader(uploaded_file)
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            return text
        
        elif file_extension in ['doc', 'docx']:
            doc = docx.Document(uploaded_file)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        
        else:
            return None
            
    except Exception as e:
        st.error(f"❌ Dosya okuma hatası: {e}")
        return None

def evaluate_with_gemini(essay_text, rubric_data):
    """Gemini AI ile essay değerlendirme"""
    
    # Rubrik kriterlerini formatla
    criteria_text = ""
    for i, criterion in enumerate(rubric_data['criteria'], 1):
        criteria_text += f"{i}. {criterion['name']} ({criterion['weight']} puan)\n"
        criteria_text += f"   Açıklama: {criterion['description']}\n"
        
        if criterion.get('levels'):
            criteria_text += "   Performans Seviyeleri:\n"
            for level, desc in criterion['levels'].items():
                if desc:
                    criteria_text += f"   - {level.title()}: {desc}\n"
        criteria_text += "\n"
    
    prompt = f"""Sen deneyimli bir öğretmensin. Aşağıdaki öğrenci yazısını verilen rubrik kriterlerine göre objektif bir şekilde değerlendir.

RUBRIK BİLGİLERİ:
- Rubrik Adı: {rubric_data['name']}
- Ders: {rubric_data.get('subject', 'Genel')}
- Toplam Puan: {rubric_data['total_points']}

DEĞERLENDIRME KRİTERLERİ:
{criteria_text}

ÖĞRENCİ METNİ:
{essay_text}

DEĞERLENDIRME FORMATI:
Lütfen aşağıdaki JSON formatında yanıt ver:

{{
    "criteria_scores": [
        {{
            "name": "Kriter Adı",
            "score": puan_sayısı,
            "max_score": maksimum_puan,
            "feedback": "Bu kritere ilişkin detaylı geri bildirim",
            "level": "mükemmel/iyi/orta/zayıf"
        }}
    ],
    "total_score": toplam_puan,
    "total_max_score": {rubric_data['total_points']},
    "percentage": yüzde_değeri,
    "grade": "harf_notu",
    "general_feedback": "Genel değerlendirme ve yorumlar",
    "strengths": ["Güçlü yön 1", "Güçlü yön 2", "Güçlü yön 3"],
    "improvements": ["Gelişim önerisi 1", "Gelişim önerisi 2", "Gelişim önerisi 3"],
    "text_statistics": {{
        "word_count": kelime_sayısı,
        "sentence_count": cümle_sayısı,
        "paragraph_count": paragraf_sayısı,
        "readability": "kolay/orta/zor"
    }}
}}

Değerlendirmeni profesyonel, yapıcı ve objektif bir dilde yap. Her kriter için verdiğin puanı gerekçelendir."""

    try:
        response = model.generate_content(prompt)
        
        # JSON'u temizle ve parse et
        response_text = response.text.strip()
        
        # JSON kısmını bul
        if "```json" in response_text:
            json_start = response_text.find("```json") + 7
            json_end = response_text.find("```", json_start)
            json_text = response_text[json_start:json_end]
        elif response_text.startswith("{"):
            json_text = response_text
        else:
            # JSON bulunamadıysa, { ile başlayan kısmı bul
            json_start = response_text.find("{")
            if json_start != -1:
                json_text = response_text[json_start:]
            else:
                raise ValueError("JSON formatı bulunamadı")
        
        # JSON'u parse et
        evaluation_result = json.loads(json_text)
        return evaluation_result
        
    except json.JSONDecodeError as e:
        st.error(f"❌ AI yanıtı JSON formatında değil: {e}")
        return None
    except Exception as e:
        st.error(f"❌ AI değerlendirme hatası: {e}")
        return None

def save_evaluation_to_db(db, evaluation_data):
    """Değerlendirmeyi veritabanına kaydet"""
    try:
        # Tarih objelerini kontrol et ve düzelt
        if 'assignment_date' in evaluation_data and evaluation_data['assignment_date']:
            from datetime import datetime, date
            if isinstance(evaluation_data['assignment_date'], date):
                # date objesi ise datetime'a çevir
                evaluation_data['assignment_date'] = datetime.combine(
                    evaluation_data['assignment_date'], 
                    datetime.min.time()
                )
        
        # MongoDB'ye kaydet
        result = db.evaluations.insert_one(evaluation_data)
        return result.inserted_id
        
    except Exception as e:
        st.error(f"❌ Veritabanına kaydetme hatası: {e}")
        return None

def grade_converter(percentage):
    """Yüzdeyi harf notuna çevir"""
    if percentage >= 90:
        return "AA"
    elif percentage >= 85:
        return "BA"
    elif percentage >= 75:
        return "BB"
    elif percentage >= 65:
        return "CB"
    elif percentage >= 55:
        return "CC"
    elif percentage >= 45:
        return "DC"
    else:
        return "FF"

# show_essay_evaluation fonksiyonunu bu şekilde değiştirin:

def show_essay_evaluation(db):
    """Ödev değerlendirme sayfası"""
    st.header("📄 Ödev Değerlendirme")
    
    # Rubrik seçimi
    rubrics = list(db.rubrics.find())
    if not rubrics:
        st.warning("⚠️ Önce bir rubrik oluşturmanız gerekiyor!")
        if st.button("📋 Rubrik Oluşturmaya Git"):
            st.session_state.page = "📋 Rubrik Yönetimi"
            st.rerun()
        return
    
    # Rubrik seçim kutusu
    rubric_options = {}
    for r in rubrics:
        display_name = f"{r['name']} ({len(r['criteria'])} kriter, {r['total_points']} puan)"
        rubric_options[display_name] = r['_id']
    
    selected_rubric_name = st.selectbox("📋 Değerlendirme Rubriği Seçin:", list(rubric_options.keys()))
    selected_rubric_id = rubric_options[selected_rubric_name]
    
    # Seçilen rubriği al
    selected_rubric = db.rubrics.find_one({"_id": selected_rubric_id})
    
    if selected_rubric:
        # Rubrik önizlemesi
        with st.expander("📋 Seçilen Rubrik Detayları", expanded=False):
            col1, col2, col3 = st.columns(3)
            with col1:
                st.metric("📚 Ders", selected_rubric.get('subject', 'Genel'))
            with col2:
                st.metric("💯 Toplam Puan", selected_rubric['total_points'])
            with col3:
                st.metric("📊 Kriter Sayısı", len(selected_rubric['criteria']))
            
            st.markdown("**📊 Değerlendirme Kriterleri:**")
            for i, criterion in enumerate(selected_rubric['criteria'], 1):
                st.write(f"{i}. **{criterion['name']}** ({criterion['weight']} puan) - {criterion['description']}")
    
    st.markdown("---")
    
    # Dosya yükleme bölümü
    st.subheader("📁 Ödev Dosyalarını Yükle")
    
    col1, col2 = st.columns([2, 1])
    
    with col1:
        uploaded_files = st.file_uploader(
            "📄 Dosya seçin (PDF, DOC, DOCX, TXT)",
            type=['pdf', 'doc', 'docx', 'txt'],
            accept_multiple_files=True,
            help="Birden fazla dosya seçebilirsiniz. Desteklenen formatlar: PDF, DOC, DOCX, TXT"
        )
    
    with col2:
        if uploaded_files:
            st.success(f"✅ {len(uploaded_files)} dosya yüklendi")
            
            total_size = sum(file.size for file in uploaded_files)
            st.info(f"📊 Toplam boyut: {total_size/1024:.1f} KB")
    
    # Yüklenen dosyaları göster
    if uploaded_files:
        st.markdown("### 📋 Yüklenen Dosyalar")
        
        for i, file in enumerate(uploaded_files, 1):
            with st.expander(f"📄 {i}. {file.name} ({file.size/1024:.1f} KB)", expanded=False):
                
                # Dosya önizlemesi
                text_content = extract_text_from_file(file)
                
                if text_content:
                    # Metin istatistikleri
                    word_count = len(text_content.split())
                    char_count = len(text_content)
                    line_count = len(text_content.split('\n'))
                    
                    col1, col2, col3 = st.columns(3)
                    with col1:
                        st.metric("📝 Kelime", word_count)
                    with col2:
                        st.metric("🔤 Karakter", char_count)
                    with col3:
                        st.metric("📄 Satır", line_count)
                    
                    # Metin önizlemesi
                    st.markdown("**📖 Metin Önizlemesi:**")
                    preview_text = text_content[:500] + "..." if len(text_content) > 500 else text_content
                    st.text_area("", value=preview_text, height=100, disabled=True, key=f"preview_{i}")
                    
                else:
                    st.error("❌ Dosya içeriği okunamadı!")
        
        st.markdown("---")
        
        # Değerlendirme seçenekleri
        st.subheader("🎯 Değerlendirme Seçenekleri")
        
        col1, col2 = st.columns(2)
        
        with col1:
            # Öğrenci bilgileri
            student_name = st.text_input("👤 Öğrenci Adı (Opsiyonel)", placeholder="Ahmet Mehmet")
            student_number = st.text_input("🆔 Öğrenci Numarası (Opsiyonel)", placeholder="12345")
        
        with col2:
            assignment_title = st.text_input("📝 Ödev Başlığı (Opsiyonel)", placeholder="Kompozisyon Ödevi")
            assignment_date = st.date_input("📅 Ödev Tarihi", value=datetime.now().date())
        
        # Ana değerlendirme butonu
        if st.button("🚀 Değerlendirmeyi Başlat", type="primary", use_container_width=True):
            
            # Progress bar
            progress_bar = st.progress(0)
            status_text = st.empty()
            
            results = []
            
            for i, file in enumerate(uploaded_files):
                status_text.text(f"📄 {file.name} değerlendiriliyor... ({i+1}/{len(uploaded_files)})")
                progress_bar.progress((i) / len(uploaded_files))
                
                # Dosya içeriğini çıkar
                text_content = extract_text_from_file(file)
                
                if text_content:
                    # AI ile değerlendir
                    status_text.text(f"🤖 AI analizi yapılıyor... ({i+1}/{len(uploaded_files)})")
                    
                    evaluation_result = evaluate_with_gemini(text_content, selected_rubric)
                    
                    if evaluation_result:
                        # Tarih objesini datetime'a çevir
                        formatted_date = datetime.combine(assignment_date, datetime.min.time()) if assignment_date else datetime.now()
                        
                        # Sonucu kaydet
                        evaluation_data = {
                            "rubric_id": selected_rubric_id,
                            "rubric_name": selected_rubric['name'],
                            "file_name": file.name,
                            "student_name": student_name if student_name else "Anonim",
                            "student_number": student_number if student_number else None,
                            "assignment_title": assignment_title if assignment_title else None,
                            "assignment_date": formatted_date,
                            "essay_text": text_content,
                            "evaluation_result": evaluation_result,
                            "total_score": evaluation_result.get('total_score', 0),
                            "percentage": evaluation_result.get('percentage', 0),
                            "grade": evaluation_result.get('grade', 'N/A'),
                            "created_at": datetime.now(),
                            "updated_at": datetime.now()
                        }
                        
                        # Veritabanına kaydet
                        evaluation_id = save_evaluation_to_db(db, evaluation_data)
                        
                        if evaluation_id:
                            evaluation_data['_id'] = evaluation_id
                            results.append(evaluation_data)
                            st.success(f"✅ {file.name} değerlendirmesi kaydedildi!")
                        else:
                            st.error(f"❌ {file.name} değerlendirmesi kaydedilemedi!")
                    else:
                        st.error(f"❌ {file.name} AI değerlendirme hatası!")
                else:
                    st.error(f"❌ {file.name} dosyası okunamadı!")
            
            # Progress tamamla
            progress_bar.progress(1.0)
            status_text.text("✅ Tüm değerlendirmeler tamamlandı!")
            
            # Sonuçları göster
            if results:
                st.success(f"🎉 {len(results)} dosya başarıyla değerlendirildi!")
                st.balloons()
                
                # Sonuçları görüntüle
                show_evaluation_results(results)
            else:
                st.error("❌ Hiçbir dosya değerlendirilemedi!")

def show_evaluation_results(results):
    """Değerlendirme sonuçlarını görüntüle"""
    st.markdown("---")
    st.header("📊 Değerlendirme Sonuçları")
    
    # Özet istatistikler
    if len(results) > 1:
        col1, col2, col3, col4 = st.columns(4)
        
        total_scores = [r['total_score'] for r in results]
        percentages = [r['percentage'] for r in results]
        
        with col1:
            st.metric("📄 Toplam Dosya", len(results))
        with col2:
            st.metric("📈 Ortalama Puan", f"{sum(total_scores)/len(total_scores):.1f}")
        with col3:
            st.metric("📊 Ortalama %", f"{sum(percentages)/len(percentages):.1f}%")
        with col4:
            st.metric("🏆 En Yüksek", f"{max(total_scores):.1f}")
    
    # Her dosya için detaylı sonuçlar
    for i, result in enumerate(results, 1):
        with st.expander(f"📄 {i}. {result['file_name']} - {result['grade']} ({result['percentage']:.1f}%)", expanded=True):
            
            # Temel bilgiler
            col1, col2, col3 = st.columns(3)
            with col1:
                st.metric("👤 Öğrenci", result['student_name'])
                if result['student_number']:
                    st.write(f"🆔 **Numara:** {result['student_number']}")
            with col2:
                st.metric("💯 Puan", f"{result['total_score']}/{result['evaluation_result']['total_max_score']}")
                st.metric("📊 Yüzde", f"{result['percentage']:.1f}%")
            with col3:
                st.metric("📝 Harf Notu", result['grade'])
                grade_color = {"AA": "🟢", "BA": "🟢", "BB": "🟡", "CB": "🟡", "CC": "🟠", "DC": "🔴", "FF": "🔴"}
                st.write(f"{grade_color.get(result['grade'], '⚪')} **Başarı Durumu**")
            
            st.markdown("---")
            
            # Kriter bazlı sonuçlar
            st.markdown("### 📊 Kriter Bazlı Değerlendirme")
            
            for criterion in result['evaluation_result']['criteria_scores']:
                progress_value = criterion['score'] / criterion['max_score']
                
                col1, col2 = st.columns([3, 1])
                with col1:
                    st.markdown(f"**{criterion['name']}**")
                    st.progress(progress_value)
                    st.write(criterion['feedback'])
                with col2:
                    st.metric("Puan", f"{criterion['score']}/{criterion['max_score']}")
                    if criterion.get('level'):
                        level_emoji = {"mükemmel": "🏆", "iyi": "👍", "orta": "👌", "zayıf": "⚠️"}
                        st.write(f"{level_emoji.get(criterion['level'], '⚪')} {criterion['level'].title()}")
            
            st.markdown("---")
            
            # Genel değerlendirme
            col1, col2 = st.columns(2)
            
            with col1:
                st.markdown("### 💪 Güçlü Yönler")
                for strength in result['evaluation_result']['strengths']:
                    st.write(f"✅ {strength}")
            
            with col2:
                st.markdown("### 📈 Gelişim Önerileri")
                for improvement in result['evaluation_result']['improvements']:
                    st.write(f"🎯 {improvement}")
            
            # Metin istatistikleri
            if result['evaluation_result'].get('text_statistics'):
                st.markdown("### 📊 Metin İstatistikleri")
                stats = result['evaluation_result']['text_statistics']
                
                col1, col2, col3, col4 = st.columns(4)
                with col1:
                    st.metric("📝 Kelime", stats.get('word_count', 'N/A'))
                with col2:
                    st.metric("📄 Cümle", stats.get('sentence_count', 'N/A'))
                with col3:
                    st.metric("📋 Paragraf", stats.get('paragraph_count', 'N/A'))
                with col4:
                    st.metric("📖 Okunabilirlik", stats.get('readability', 'N/A'))
            
            # Genel feedback
            st.markdown("### 📝 Genel Değerlendirme")
            st.info(result['evaluation_result']['general_feedback'])
            
            # Eylem butonları
            st.markdown("---")
            col1, col2, col3 = st.columns(3)
            
            with col1:
                if st.button(f"📄 PDF Rapor İndir", key=f"pdf_{i}"):
                    st.info("PDF rapor özelliği yakında eklenecek!")
            
            with col2:
                if st.button(f"📧 E-posta Gönder", key=f"email_{i}"):
                    st.info("E-posta gönderme özelliği yakında eklenecek!")
            
            with col3:
                if st.button(f"🔄 Yeniden Değerlendir", key=f"reeval_{i}"):
                    st.info("Yeniden değerlendirme özelliği yakında eklenecek!")

def extract_text_from_file(uploaded_file):
    """Dosyadan metin çıkarma"""
    try:
        file_extension = uploaded_file.name.split('.')[-1].lower()
        
        if file_extension == 'txt':
            return uploaded_file.read().decode('utf-8')
        
        elif file_extension == 'pdf':
            pdf_reader = PyPDF2.PdfReader(uploaded_file)
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            return text
        
        elif file_extension in ['doc', 'docx']:
            doc = docx.Document(uploaded_file)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        
        else:
            return None
            
    except Exception as e:
        st.error(f"❌ Dosya okuma hatası: {e}")
        return None

def evaluate_with_gemini(essay_text, rubric_data):
    """Gemini AI ile essay değerlendirme"""
    
    # Rubrik kriterlerini formatla
    criteria_text = ""
    for i, criterion in enumerate(rubric_data['criteria'], 1):
        criteria_text += f"{i}. {criterion['name']} ({criterion['weight']} puan)\n"
        criteria_text += f"   Açıklama: {criterion['description']}\n"
        
        if criterion.get('levels'):
            criteria_text += "   Performans Seviyeleri:\n"
            for level, desc in criterion['levels'].items():
                if desc:
                    criteria_text += f"   - {level.title()}: {desc}\n"
        criteria_text += "\n"
    
    prompt = f"""Sen deneyimli bir öğretmensin. Aşağıdaki öğrenci yazısını verilen rubrik kriterlerine göre objektif bir şekilde değerlendir.

RUBRIK BİLGİLERİ:
- Rubrik Adı: {rubric_data['name']}
- Ders: {rubric_data.get('subject', 'Genel')}
- Toplam Puan: {rubric_data['total_points']}

DEĞERLENDIRME KRİTERLERİ:
{criteria_text}

ÖĞRENCİ METNİ:
{essay_text}

DEĞERLENDIRME FORMATI:
Lütfen aşağıdaki JSON formatında yanıt ver:

{{
    "criteria_scores": [
        {{
            "name": "Kriter Adı",
            "score": puan_sayısı,
            "max_score": maksimum_puan,
            "feedback": "Bu kritere ilişkin detaylı geri bildirim",
            "level": "mükemmel/iyi/orta/zayıf"
        }}
    ],
    "total_score": toplam_puan,
    "total_max_score": {rubric_data['total_points']},
    "percentage": yüzde_değeri,
    "grade": "harf_notu",
    "general_feedback": "Genel değerlendirme ve yorumlar",
    "strengths": ["Güçlü yön 1", "Güçlü yön 2", "Güçlü yön 3"],
    "improvements": ["Gelişim önerisi 1", "Gelişim önerisi 2", "Gelişim önerisi 3"],
    "text_statistics": {{
        "word_count": kelime_sayısı,
        "sentence_count": cümle_sayısı,
        "paragraph_count": paragraf_sayısı,
        "readability": "kolay/orta/zor"
    }}
}}

Değerlendirmeni profesyonel, yapıcı ve objektif bir dilde yap. Her kriter için verdiğin puanı gerekçelendir."""

    try:
        response = model.generate_content(prompt)
        
        # JSON'u temizle ve parse et
        response_text = response.text.strip()
        
        # JSON kısmını bul
        if "```json" in response_text:
            json_start = response_text.find("```json") + 7
            json_end = response_text.find("```", json_start)
            json_text = response_text[json_start:json_end]
        elif response_text.startswith("{"):
            json_text = response_text
        else:
            # JSON bulunamadıysa, { ile başlayan kısmı bul
            json_start = response_text.find("{")
            if json_start != -1:
                json_text = response_text[json_start:]
            else:
                raise ValueError("JSON formatı bulunamadı")
        
        # JSON'u parse et
        evaluation_result = json.loads(json_text)
        
        # Harf notunu hesapla
        percentage = evaluation_result.get('percentage', 0)
        evaluation_result['grade'] = grade_converter(percentage)
        
        return evaluation_result
        
    except json.JSONDecodeError as e:
        st.error(f"❌ AI yanıtı JSON formatında değil: {e}")
        return None
    except Exception as e:
        st.error(f"❌ AI değerlendirme hatası: {e}")
        return None

def grade_converter(percentage):
    """Yüzdeyi harf notuna çevir"""
    if percentage >= 90:
        return "AA"
    elif percentage >= 85:
        return "BA"
    elif percentage >= 75:
        return "BB"
    elif percentage >= 65:
        return "CB"
    elif percentage >= 55:
        return "CC"
    elif percentage >= 45:
        return "DC"
    else:
        return "FF"


# Ana uygulama
def main():
    st.title("📝 Essay Grader AI")
    st.subheader("Yapay Zeka ile Otomatik Ödev Değerlendirme")
    
    # MongoDB bağlantısı
    try:
        db = init_mongodb()
        
    except Exception as e:
        st.sidebar.error(f"❌ Veritabanı hatası: {e}")
        return
    
    # Sidebar - Ana menü
    st.sidebar.title("🎯 Ana Menü")
    page = st.sidebar.selectbox(
        "Sayfa Seç:",
        ["🏠 Ana Sayfa", "📋 Rubrik Yönetimi", "📄 Ödev Değerlendirme", "📊 Raporlar"]
    )
    
    if page == "🏠 Ana Sayfa":
        show_homepage(db)
    elif page == "📋 Rubrik Yönetimi":
        show_rubric_management(db)
    elif page == "📄 Ödev Değerlendirme":
        show_essay_evaluation(db)
    elif page == "📊 Raporlar":
        show_reports(db)

def show_homepage(db):
    """Ana sayfa"""
    col1, col2, col3 = st.columns(3)
    
    with col1:
        st.metric("📋 Toplam Rubrik", get_rubric_count(db))
    
    with col2:
        st.metric("📄 Değerlendirilen Ödevler", get_evaluation_count(db))
    
    with col3:
        st.metric("👥 Aktif Öğrenciler", get_student_count(db))
    
    st.markdown("---")
    
    st.markdown("""
    ## 🚀 Hoş Geldiniz!
    
    Essay Grader AI ile ödevlerinizi hızlı ve objektif şekilde değerlendirebilirsiniz.
    
    ### ✨ Özellikler:
    - 📋 **Hazır Rubrik Şablonları**: Kompozisyon, deneme, proje raporu
    - 🤖 **AI Değerlendirme**: Gemini ile detaylı analiz
    - 📊 **Detaylı Raporlar**: Öğrenci ve sınıf bazlı analizler
    - 📄 **Çoklu Format**: PDF, DOC, DOCX, TXT dosya desteği
    
    ### 🎯 Hızlı Başlangıç:
    1. **Rubrik Yönetimi** → Değerlendirme kriteri seç
    2. **Ödev Değerlendirme** → Dosyaları yükle
    3. **Raporlar** → Sonuçları analiz et
    """)

def show_rubric_management(db):
    """Rubrik yönetimi sayfası"""
    st.header("📋 Rubrik Yönetimi")
    
    tab1, tab2 = st.tabs(["🎯 Mevcut Rubrikler", "➕ Yeni Rubrik"])
    
    with tab1:
        # Şablon ve özel rubrikleri ayır
        template_rubrics = list(db.rubrics.find({"is_template": True}))
        custom_rubrics = list(db.rubrics.find({"is_template": False}))
        
        # Hazır Şablonlar Bölümü
        st.subheader("📚 Hazır Şablonlar")
        
        if template_rubrics:
            for template in template_rubrics:
                with st.expander(f"📋 {template['name']} ({len(template['criteria'])} kriter)", expanded=False):
                    col1, col2 = st.columns([3, 1])
                    
                    with col1:
                        st.write(f"**📝 Açıklama:** {template['description']}")
                        st.write(f"**🎓 Sınıf Seviyeleri:** {', '.join(map(str, template['grade_levels']))}")
                        st.write(f"**📚 Ders:** {template.get('subject', 'Genel')}")
                        st.write(f"**💯 Toplam Puan:** {template['total_points']}")
                        
                        st.write("**📊 Kriterler:**")
                        for i, criterion in enumerate(template['criteria'], 1):
                            st.write(f"  {i}. **{criterion['name']}** ({criterion['weight']} puan)")
                            st.write(f"     └ {criterion['description']}")
                    
                    with col2:
                        st.markdown("**İşlemler:**")
                        
                        if st.button(f"📋 Kopyala", key=f"copy_{template['_id']}", 
                                   help="Bu şablonu özel rubrik olarak kopyalar"):
                            copy_rubric_template(db, template['_id'])
                            st.success("✅ Rubrik kopyalandı!")
                            st.rerun()
                        
                        if st.button(f"👁️ Detay", key=f"detail_{template['_id']}", 
                                   help="Detaylı görünüm"):
                            st.session_state[f"show_detail_{template['_id']}"] = True
                        
                        # Detay gösterimi
                        if st.session_state.get(f"show_detail_{template['_id']}", False):
                            st.info("👆 Yukarıdaki expander'ı açın")
        else:
            st.info("📭 Henüz hazır şablon bulunamadı.")
        
        st.markdown("---")
        
        # Özel Rubrikler Bölümü
        st.subheader("🎯 Özel Rubriklerim")
        
        if custom_rubrics:
            for rubric in custom_rubrics:
                with st.expander(f"🎯 {rubric['name']} ({len(rubric['criteria'])} kriter)", expanded=False):
                    col1, col2 = st.columns([3, 1])
                    
                    with col1:
                        st.write(f"**📝 Açıklama:** {rubric['description']}")
                        st.write(f"**🎓 Sınıf Seviyeleri:** {', '.join(map(str, rubric['grade_levels']))}")
                        st.write(f"**📚 Ders:** {rubric.get('subject', 'Genel')}")
                        st.write(f"**💯 Toplam Puan:** {rubric['total_points']}")
                        st.write(f"**📅 Oluşturulma:** {rubric['created_at'].strftime('%d.%m.%Y %H:%M')}")
                        
                        st.write("**📊 Kriterler:**")
                        for i, criterion in enumerate(rubric['criteria'], 1):
                            st.write(f"  {i}. **{criterion['name']}** ({criterion['weight']} puan)")
                            st.write(f"     └ {criterion['description']}")
                            
                            # Eğer detaylı seviyeler varsa göster
                            if criterion.get('levels'):
                                st.write("     **Performans Seviyeleri:**")
                                for level, desc in criterion['levels'].items():
                                    if desc:  # Boş olmayan seviyeler
                                        level_emoji = {"mükemmel": "🏆", "iyi": "👍", "orta": "👌", "zayıf": "⚠️"}
                                        st.write(f"       {level_emoji.get(level, '•')} {level.title()}: {desc}")
                    
                    with col2:
                        st.markdown("**İşlemler:**")
                        
                        # Düzenleme butonu
                        if st.button(f"✏️ Düzenle", key=f"edit_{rubric['_id']}", 
                                   help="Bu rubriği düzenle"):
                            st.session_state[f"edit_rubric_{rubric['_id']}"] = True
                            st.info("🚧 Düzenleme özelliği yakında eklenecek!")
                        
                        # Kopyalama butonu
                        if st.button(f"📄 Kopya Oluştur", key=f"duplicate_{rubric['_id']}", 
                                   help="Bu rubrikten bir kopya oluştur"):
                            duplicate_rubric(db, rubric['_id'])
                            st.success("✅ Kopya oluşturuldu!")
                            st.rerun()
                        
                        # Silme butonu (onay ile)
                        if st.button(f"🗑️ Sil", key=f"delete_{rubric['_id']}", 
                                   help="Bu rubriği kalıcı olarak sil", 
                                   type="secondary"):
                            # Silme onayı
                            if f"confirm_delete_{rubric['_id']}" not in st.session_state:
                                st.session_state[f"confirm_delete_{rubric['_id']}"] = False
                            
                            if not st.session_state[f"confirm_delete_{rubric['_id']}"]:
                                st.warning("⚠️ Silme işlemini onaylamak için tekrar tıklayın!")
                                st.session_state[f"confirm_delete_{rubric['_id']}"] = True
                            else:
                                # Gerçek silme işlemi
                                try:
                                    result = db.rubrics.delete_one({"_id": rubric['_id']})
                                    if result.deleted_count > 0:
                                        st.success(f"✅ '{rubric['name']}' rubriği silindi!")
                                        # Session state temizle
                                        if f"confirm_delete_{rubric['_id']}" in st.session_state:
                                            del st.session_state[f"confirm_delete_{rubric['_id']}"]
                                        st.rerun()
                                    else:
                                        st.error("❌ Rubrik silinirken hata oluştu!")
                                except Exception as e:
                                    st.error(f"❌ Silme hatası: {e}")
                        
                        # İstatistikler
                        st.markdown("**📈 İstatistikler:**")
                        
                        # Bu rubrikle yapılan değerlendirme sayısı
                        evaluation_count = db.evaluations.count_documents({"rubric_id": rubric['_id']})
                        st.write(f"📊 Değerlendirme: {evaluation_count}")
                        
                        # Ortalama puan (eğer değerlendirme varsa)
                        if evaluation_count > 0:
                            avg_score = db.evaluations.aggregate([
                                {"$match": {"rubric_id": rubric['_id']}},
                                {"$group": {"_id": None, "avg_score": {"$avg": "$total_score"}}}
                            ])
                            avg_result = list(avg_score)
                            if avg_result:
                                st.write(f"📈 Ort. Puan: {avg_result[0]['avg_score']:.1f}")
        else:
            st.info("📭 Henüz özel rubrik oluşturmamışsınız.")
            st.markdown("**💡 İpucu:** Yukarıdaki şablonları kopyalayarak başlayabilir veya sağdaki sekmeden yeni rubrik oluşturabilirsiniz.")
        
        # Genel istatistikler
        st.markdown("---")
        st.subheader("📊 Genel İstatistikler")
        
        col1, col2, col3, col4 = st.columns(4)
        
        with col1:
            st.metric("📚 Toplam Şablon", len(template_rubrics))
        
        with col2:
            st.metric("🎯 Özel Rubrik", len(custom_rubrics))
        
        with col3:
            total_evaluations = db.evaluations.count_documents({})
            st.metric("📊 Toplam Değerlendirme", total_evaluations)
        
        with col4:
            active_rubrics = len([r for r in custom_rubrics + template_rubrics if db.evaluations.count_documents({"rubric_id": r['_id']}) > 0])
            st.metric("✅ Aktif Rubrik", active_rubrics)
    
    with tab2:
        st.subheader("➕ Yeni Rubrik Oluştur")
        create_new_rubric_form(db)

def duplicate_rubric(db, rubric_id):
    """Mevcut rubrikten kopya oluşturur"""
    try:
        original = db.rubrics.find_one({"_id": rubric_id})
        if original:
            new_rubric = original.copy()
            new_rubric.pop('_id')
            new_rubric['name'] = f"{original['name']} - Kopya"
            new_rubric['is_template'] = False
            new_rubric['created_at'] = datetime.now()
            new_rubric['updated_at'] = datetime.now()
            
            result = db.rubrics.insert_one(new_rubric)
            return result.inserted_id
    except Exception as e:
        st.error(f"Kopyalama hatası: {e}")
        return None

def copy_rubric_template(db, template_id):
    """Şablon rubriği özel rubrik olarak kopyalar"""
    try:
        template = db.rubrics.find_one({"_id": template_id})
        if template:
            new_rubric = template.copy()
            new_rubric.pop('_id')
            new_rubric['name'] = f"{template['name']} - Özel"
            new_rubric['is_template'] = False
            new_rubric['teacher_id'] = None  # Sonra öğretmen sistemi eklenecek
            new_rubric['created_at'] = datetime.now()
            new_rubric['updated_at'] = datetime.now()
            
            result = db.rubrics.insert_one(new_rubric)
            return result.inserted_id
    except Exception as e:
        st.error(f"Rubrik kopyalama hatası: {e}")
        return None

def show_evaluation_results(results):
    """Değerlendirme sonuçlarını görüntüle"""
    st.markdown("---")
    st.header("📊 Değerlendirme Sonuçları")
    
    # Özet istatistikler
    if len(results) > 1:
        col1, col2, col3, col4 = st.columns(4)
        
        total_scores = [r['total_score'] for r in results]
        percentages = [r['percentage'] for r in results]
        
        with col1:
            st.metric("📄 Toplam Dosya", len(results))
        with col2:
            st.metric("📈 Ortalama Puan", f"{sum(total_scores)/len(total_scores):.1f}")
        with col3:
            st.metric("📊 Ortalama %", f"{sum(percentages)/len(percentages):.1f}%")
        with col4:
            st.metric("🏆 En Yüksek", f"{max(total_scores):.1f}")
    
    # Her dosya için detaylı sonuçlar
    for i, result in enumerate(results, 1):
        with st.expander(f"📄 {i}. {result['file_name']} - {result['grade']} ({result['percentage']:.1f}%)", expanded=True):
            
            # Temel bilgiler
            col1, col2, col3 = st.columns(3)
            with col1:
                st.metric("👤 Öğrenci", result['student_name'])
                if result['student_number']:
                    st.write(f"🆔 **Numara:** {result['student_number']}")
            with col2:
                st.metric("💯 Puan", f"{result['total_score']}/{result['evaluation_result']['total_max_score']}")
                st.metric("📊 Yüzde", f"{result['percentage']:.1f}%")
            with col3:
                st.metric("📝 Harf Notu", result['grade'])
                grade_color = {"AA": "🟢", "BA": "🟢", "BB": "🟡", "CB": "🟡", "CC": "🟠", "DC": "🔴", "FF": "🔴"}
                st.write(f"{grade_color.get(result['grade'], '⚪')} **Başarı Durumu**")
            
            st.markdown("---")
            
            # Kriter bazlı sonuçlar
            st.markdown("### 📊 Kriter Bazlı Değerlendirme")
            
            for criterion in result['evaluation_result']['criteria_scores']:
                progress_value = criterion['score'] / criterion['max_score']
                
                col1, col2 = st.columns([3, 1])
                with col1:
                    st.markdown(f"**{criterion['name']}**")
                    st.progress(progress_value)
                    st.write(criterion['feedback'])
                with col2:
                    st.metric("Puan", f"{criterion['score']}/{criterion['max_score']}")
                    if criterion.get('level'):
                        level_emoji = {"mükemmel": "🏆", "iyi": "👍", "orta": "👌", "zayıf": "⚠️"}
                        st.write(f"{level_emoji.get(criterion['level'], '⚪')} {criterion['level'].title()}")
            
            st.markdown("---")
            
            # Genel değerlendirme
            col1, col2 = st.columns(2)
            
            with col1:
                st.markdown("### 💪 Güçlü Yönler")
                for strength in result['evaluation_result']['strengths']:
                    st.write(f"✅ {strength}")
            
            with col2:
                st.markdown("### 📈 Gelişim Önerileri")
                for improvement in result['evaluation_result']['improvements']:
                    st.write(f"🎯 {improvement}")
            
            # Metin istatistikleri
            if result['evaluation_result'].get('text_statistics'):
                st.markdown("### 📊 Metin İstatistikleri")
                stats = result['evaluation_result']['text_statistics']
                
                col1, col2, col3, col4 = st.columns(4)
                with col1:
                    st.metric("📝 Kelime", stats.get('word_count', 'N/A'))
                with col2:
                    st.metric("📄 Cümle", stats.get('sentence_count', 'N/A'))
                with col3:
                    st.metric("📋 Paragraf", stats.get('paragraph_count', 'N/A'))
                with col4:
                    st.metric("📖 Okunabilirlik", stats.get('readability', 'N/A'))
            
            # Genel feedback
            st.markdown("### 📝 Genel Değerlendirme")
            st.info(result['evaluation_result']['general_feedback'])
            
            # Eylem butonları
            st.markdown("---")
            col1, col2, col3 = st.columns(3)
            
            with col1:
                if st.button(f"📄 PDF Rapor İndir", key=f"pdf_{i}"):
                    st.info("PDF rapor özelliği yakında eklenecek!")
            
            with col2:
                if st.button(f"📧 E-posta Gönder", key=f"email_{i}"):
                    st.info("E-posta gönderme özelliği yakında eklenecek!")
            
            with col3:
                if st.button(f"🔄 Yeniden Değerlendir", key=f"reeval_{i}"):
                    st.info("Yeniden değerlendirme özelliği yakında eklenecek!")

def show_reports(db):
    """Raporlar sayfası"""
    st.header("📊 Raporlar")
    st.info("Rapor özellikleri yakında eklenecek!")

# Yardımcı fonksiyonlar
def get_rubric_count(db):
    try:
        return db.rubrics.count_documents({})
    except Exception as e:
        st.error(f"Rubrik sayısı alınırken hata: {e}")
        return 0

def get_evaluation_count(db):
    try:
        return db.evaluations.count_documents({})
    except Exception as e:
        st.error(f"Değerlendirme sayısı alınırken hata: {e}")
        return 0

def get_student_count(db):
    try:
        return db.students.count_documents({"is_active": True})
    except Exception as e:
        st.error(f"Öğrenci sayısı alınırken hata: {e}")
        return 0

def copy_rubric_template(db, template_id):
    """Şablon rubriği kopyalar"""
    try:
        template = db.rubrics.find_one({"_id": template_id})
        if template:
            new_rubric = template.copy()
            new_rubric.pop('_id')
            new_rubric['is_template'] = False
            new_rubric['teacher_id'] = None  # Sonra öğretmen sistemi eklenecek
            new_rubric['created_at'] = datetime.now()
            db.rubrics.insert_one(new_rubric)
    except Exception as e:
        st.error(f"Rubrik kopyalama hatası: {e}")

def create_new_rubric_form(db):
    """Yeni rubrik oluşturma formu"""
    st.subheader("➕ Yeni Rubrik Oluştur")
    
    with st.form("new_rubric_form"):
        # Temel bilgiler
        col1, col2 = st.columns(2)
        
        with col1:
            rubric_name = st.text_input("📋 Rubrik Adı*", placeholder="Örn: Hikaye Yazma Rubriği")
            subject = st.selectbox("📚 Ders", [
                "Türk Dili ve Edebiyatı", 
                "İngilizce", 
                "Tarih", 
                "Coğrafya", 
                "Felsefe",
                "Sosyoloji",
                "Biyoloji",
                "Genel"
            ])
        
        with col2:
            description = st.text_area("📝 Açıklama*", placeholder="Bu rubriğin kullanım amacını açıklayın...")
            grade_levels = st.multiselect("🎓 Sınıf Seviyeleri*", 
                                        options=[9, 10, 11, 12], 
                                        default=[9, 10, 11, 12])
        
        st.markdown("---")
        st.subheader("📊 Değerlendirme Kriterleri")
        
        # Dinamik kriter ekleme
        if 'criteria_count' not in st.session_state:
            st.session_state.criteria_count = 3
        
        criteria = []
        total_weight = 0
        
        for i in range(st.session_state.criteria_count):
            st.markdown(f"**Kriter {i+1}:**")
            
            crit_col1, crit_col2, crit_col3 = st.columns([2, 1, 1])
            
            with crit_col1:
                crit_name = st.text_input(f"Kriter Adı", key=f"crit_name_{i}", 
                                        placeholder="Örn: İçerik ve Konu İşleme")
                crit_desc = st.text_input(f"Açıklama", key=f"crit_desc_{i}",
                                        placeholder="Bu kriterin değerlendirme odağı...")
            
            with crit_col2:
                crit_weight = st.number_input(f"Ağırlık (Puan)", key=f"crit_weight_{i}", 
                                            min_value=1, max_value=50, value=25)
            
            with crit_col3:
                # Detaylı seviye tanımları
                show_levels = st.checkbox(f"Detaylı Seviyeler", key=f"show_levels_{i}")
            
            # Kriter seviye tanımları
            levels = {}
            if show_levels:
                st.markdown("**Performans Seviyeleri:**")
                level_col1, level_col2 = st.columns(2)
                
                with level_col1:
                    levels["mükemmel"] = st.text_input(f"Mükemmel ({int(crit_weight*0.9)}-{crit_weight} puan)", 
                                                    key=f"excellent_{i}",
                                                    placeholder="En yüksek performans tanımı...")
                    levels["orta"] = st.text_input(f"Orta ({int(crit_weight*0.5)}-{int(crit_weight*0.7)} puan)", 
                                                 key=f"average_{i}",
                                                 placeholder="Ortalama performans tanımı...")
                
                with level_col2:
                    levels["iyi"] = st.text_input(f"İyi ({int(crit_weight*0.7)}-{int(crit_weight*0.9)} puan)", 
                                                key=f"good_{i}",
                                                placeholder="İyi performans tanımı...")
                    levels["zayıf"] = st.text_input(f"Zayıf (0-{int(crit_weight*0.5)} puan)", 
                                                  key=f"poor_{i}",
                                                  placeholder="Düşük performans tanımı...")
            
            if crit_name and crit_desc:
                criteria.append({
                    "name": crit_name,
                    "description": crit_desc,
                    "weight": crit_weight,
                    "max_points": crit_weight,
                    "levels": levels if show_levels else {}
                })
                total_weight += crit_weight
            
            st.markdown("---")
        
        # Kriter ekleme/çıkarma butonları
        col1, col2, col3 = st.columns([1, 1, 2])
        
        with col1:
            if st.form_submit_button("➕ Kriter Ekle"):
                st.session_state.criteria_count += 1
                st.rerun()
        
        with col2:
            if st.form_submit_button("➖ Kriter Çıkar") and st.session_state.criteria_count > 1:
                st.session_state.criteria_count -= 1
                st.rerun()
        
        with col3:
            st.info(f"💯 Toplam Puan: {total_weight}")
        
        # Form gönderimi
        submitted = st.form_submit_button("🚀 Rubriği Kaydet", type="primary")
        
        if submitted:
            # Validasyon
            if not rubric_name:
                st.error("❌ Rubrik adı zorunludur!")
                return
            
            if not description:
                st.error("❌ Açıklama zorunludur!")
                return
            
            if not grade_levels:
                st.error("❌ En az bir sınıf seviyesi seçilmelidir!")
                return
            
            if not criteria:
                st.error("❌ En az bir kriter tanımlanmalıdır!")
                return
            
            if total_weight == 0:
                st.error("❌ Toplam puan sıfırdan büyük olmalıdır!")
                return
            
            # Rubriği kaydet
            try:
                new_rubric = {
                    "name": rubric_name,
                    "description": description,
                    "grade_levels": grade_levels,
                    "subject": subject,
                    "criteria": criteria,
                    "total_points": total_weight,
                    "is_template": False,  # Öğretmen tarafından oluşturulan
                    "teacher_id": None,  # Sonra öğretmen sistemi eklenecek
                    "created_at": datetime.now(),
                    "updated_at": datetime.now()
                }
                
                result = db.rubrics.insert_one(new_rubric)
                
                st.success(f"✅ '{rubric_name}' rubriği başarıyla oluşturuldu!")
                st.balloons()
                
                # Session state'i temizle
                if 'criteria_count' in st.session_state:
                    del st.session_state.criteria_count
                
                # Formu sıfırla
                st.rerun()
                
            except Exception as e:
                st.error(f"❌ Rubrik kaydedilirken hata oluştu: {e}")

def show_rubric_detail(rubric):
    """Rubrik detaylarını güzel bir şekilde gösterir"""
    st.markdown(f"### 📋 {rubric['name']}")
    
    col1, col2, col3 = st.columns(3)
    with col1:
        st.metric("📚 Ders", rubric['subject'])
    with col2:
        st.metric("💯 Toplam Puan", rubric['total_points'])
    with col3:
        st.metric("🎓 Sınıf Seviyeleri", f"{len(rubric['grade_levels'])} seviye")
    
    st.markdown(f"**📝 Açıklama:** {rubric['description']}")
    st.markdown(f"**🎓 Sınıflar:** {', '.join(map(str, rubric['grade_levels']))}")
    
    st.markdown("---")
    st.markdown("### 📊 Değerlendirme Kriterleri")
    
    for i, criterion in enumerate(rubric['criteria'], 1):
        with st.expander(f"{i}. {criterion['name']} ({criterion['weight']} puan)"):
            st.write(f"**Açıklama:** {criterion['description']}")
            
            if criterion.get('levels'):
                st.write("**Performans Seviyeleri:**")
                for level, desc in criterion['levels'].items():
                    if desc:  # Boş olmayan seviyeler
                        level_emoji = {"mükemmel": "🏆", "iyi": "👍", "orta": "👌", "zayıf": "⚠️"}
                        st.write(f"  {level_emoji.get(level, '•')} **{level.title()}:** {desc}")

def show_reports(db):
    """Raporlar sayfası"""
    st.header("📊 Raporlar ve Analizler")
    
    # Genel istatistikleri al
    total_evaluations = db.evaluations.count_documents({})
    total_students = len(db.evaluations.distinct("student_name"))
    total_rubrics = db.rubrics.count_documents({})
    
    if total_evaluations == 0:
        st.warning("📭 Henüz değerlendirme yapılmamış. Önce ödev değerlendirme sayfasından dosya analizi yapın.")
        if st.button("📄 Ödev Değerlendirme Sayfasına Git"):
            st.session_state.page = "📄 Ödev Değerlendirme"
            st.rerun()
        return
    
    # Ana dashboard istatistikleri
    st.subheader("📈 Genel İstatistikler")
    
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        st.metric("📄 Toplam Değerlendirme", total_evaluations)
    
    with col2:
        st.metric("👥 Öğrenci Sayısı", total_students)
    
    with col3:
        # Ortalama puan hesapla
        avg_pipeline = [{"$group": {"_id": None, "avg_score": {"$avg": "$percentage"}}}]
        avg_result = list(db.evaluations.aggregate(avg_pipeline))
        avg_score = avg_result[0]['avg_score'] if avg_result else 0
        st.metric("📊 Ortalama Başarı", f"{avg_score:.1f}%")
    
    with col4:
        # Başarılı öğrenci oranı (60% üzeri)
        success_count = db.evaluations.count_documents({"percentage": {"$gte": 60}})
        success_rate = (success_count / total_evaluations) * 100 if total_evaluations > 0 else 0
        st.metric("🏆 Başarı Oranı", f"{success_rate:.1f}%")
    
    st.markdown("---")
    
    # Tab'lar ile farklı rapor türleri
    tab1, tab2, tab3, tab4, tab5 = st.tabs([
        "📊 Dashboard", 
        "👥 Öğrenci Raporları", 
        "📋 Rubrik Analizleri", 
        "📈 Grafik Analizler", 
        "📄 Export"
    ])
    
    with tab1:
        show_dashboard_report(db)
    
    with tab2:
        show_student_reports(db)
    
    with tab3:
        show_rubric_analysis(db)
    
    with tab4:
        show_graphic_analysis(db)
    
    with tab5:
        show_export_options(db)

def show_dashboard_report(db):
    """Dashboard ana rapor"""
    st.subheader("📈 Dashboard Özeti")
    
    # Son değerlendirmeler
    st.markdown("### 🕒 Son Değerlendirmeler")
    
    recent_evaluations = list(db.evaluations.find().sort("created_at", -1).limit(5))
    
    if recent_evaluations:
        for eval_data in recent_evaluations:
            col1, col2, col3, col4 = st.columns([2, 1, 1, 1])
            
            with col1:
                st.write(f"📄 **{eval_data['file_name']}**")
                st.write(f"👤 {eval_data['student_name']}")
            
            with col2:
                st.write(f"📋 {eval_data['rubric_name']}")
            
            with col3:
                grade_color = {"AA": "🟢", "BA": "🟢", "BB": "🟡", "CB": "🟡", "CC": "🟠", "DC": "🔴", "FF": "🔴"}
                st.write(f"{grade_color.get(eval_data['grade'], '⚪')} **{eval_data['grade']}**")
                st.write(f"{eval_data['percentage']:.1f}%")
            
            with col4:
                st.write(f"📅 {eval_data['created_at'].strftime('%d.%m.%Y')}")
                st.write(f"🕐 {eval_data['created_at'].strftime('%H:%M')}")
    
    st.markdown("---")
    
    # Harf notu dağılımı
    st.markdown("### 📊 Harf Notu Dağılımı")
    
    grade_pipeline = [
        {"$group": {"_id": "$grade", "count": {"$sum": 1}}},
        {"$sort": {"_id": 1}}
    ]
    grade_distribution = list(db.evaluations.aggregate(grade_pipeline))
    
    if grade_distribution:
        col1, col2 = st.columns(2)
        
        with col1:
            # Tablo görünümü
            for grade_data in grade_distribution:
                grade = grade_data['_id']
                count = grade_data['count']
                percentage = (count / sum(g['count'] for g in grade_distribution)) * 100
                
                grade_color = {"AA": "🟢", "BA": "🟢", "BB": "🟡", "CB": "🟡", "CC": "🟠", "DC": "🔴", "FF": "🔴"}
                st.write(f"{grade_color.get(grade, '⚪')} **{grade}**: {count} öğrenci ({percentage:.1f}%)")
        
        with col2:
            # Basit progress bar görünümü
            total_count = sum(g['count'] for g in grade_distribution)
            st.write("**Görsel Dağılım:**")
            
            for grade_data in grade_distribution:
                grade = grade_data['_id']
                count = grade_data['count']
                progress_value = count / total_count
                st.write(f"**{grade}**")
                st.progress(progress_value)
    
    st.markdown("---")
    
    # En çok kullanılan rubrikler
    st.markdown("### 📋 Popüler Rubrikler")
    
    rubric_pipeline = [
        {"$group": {"_id": "$rubric_name", "count": {"$sum": 1}, "avg_score": {"$avg": "$percentage"}}},
        {"$sort": {"count": -1}},
        {"$limit": 5}
    ]
    popular_rubrics = list(db.evaluations.aggregate(rubric_pipeline))
    
    if popular_rubrics:
        for rubric_data in popular_rubrics:
            col1, col2, col3 = st.columns([2, 1, 1])
            
            with col1:
                st.write(f"📋 **{rubric_data['_id']}**")
            
            with col2:
                st.metric("📊 Kullanım", rubric_data['count'])
            
            with col3:
                st.metric("📈 Ortalama", f"{rubric_data['avg_score']:.1f}%")

def show_student_reports(db):
    """Öğrenci bazlı raporlar"""
    st.subheader("👥 Öğrenci Raporları")
    
    # Öğrenci listesi
    students = db.evaluations.distinct("student_name")
    students = [s for s in students if s and s != "Anonim"]
    
    if not students:
        st.info("📭 Adı girilmiş öğrenci bulunamadı. Değerlendirme yaparken öğrenci adı girmeyi unutmayın.")
        return
    
    # Öğrenci seçimi
    selected_student = st.selectbox("👤 Öğrenci Seçin:", ["Tümü"] + students)
    
    if selected_student == "Tümü":
        # Tüm öğrencilerin özet raporu
        st.markdown("### 📊 Tüm Öğrenciler Özet Rapor")
        
        student_stats_pipeline = [
            {"$match": {"student_name": {"$ne": "Anonim"}}},
            {"$group": {
                "_id": "$student_name",
                "total_evaluations": {"$sum": 1},
                "avg_score": {"$avg": "$percentage"},
                "max_score": {"$max": "$percentage"},
                "min_score": {"$min": "$percentage"},
                "latest_date": {"$max": "$created_at"}
            }},
            {"$sort": {"avg_score": -1}}
        ]
        
        student_stats = list(db.evaluations.aggregate(student_stats_pipeline))
        
        if student_stats:
            # Tablo başlığı
            col1, col2, col3, col4, col5, col6 = st.columns([2, 1, 1, 1, 1, 1])
            
            with col1:
                st.write("**👤 Öğrenci Adı**")
            with col2:
                st.write("**📊 Değerlendirme**")
            with col3:
                st.write("**📈 Ortalama**")
            with col4:
                st.write("**🏆 En Yüksek**")
            with col5:
                st.write("**📉 En Düşük**")
            with col6:
                st.write("**📅 Son Tarih**")
            
            st.markdown("---")
            
            # Öğrenci verileri
            for student_data in student_stats:
                col1, col2, col3, col4, col5, col6 = st.columns([2, 1, 1, 1, 1, 1])
                
                with col1:
                    st.write(f"👤 {student_data['_id']}")
                
                with col2:
                    st.write(f"{student_data['total_evaluations']} adet")
                
                with col3:
                    avg_score = student_data['avg_score']
                    grade = grade_converter(avg_score)
                    grade_color = {"AA": "🟢", "BA": "🟢", "BB": "🟡", "CB": "🟡", "CC": "🟠", "DC": "🔴", "FF": "🔴"}
                    st.write(f"{grade_color.get(grade, '⚪')} {avg_score:.1f}%")
                
                with col4:
                    st.write(f"🏆 {student_data['max_score']:.1f}%")
                
                with col5:
                    st.write(f"📉 {student_data['min_score']:.1f}%")
                
                with col6:
                    st.write(f"📅 {student_data['latest_date'].strftime('%d.%m.%Y')}")
    
    else:
        # Seçili öğrencinin detaylı raporu
        st.markdown(f"### 👤 {selected_student} - Detaylı Rapor")
        
        student_evaluations = list(db.evaluations.find({"student_name": selected_student}).sort("created_at", -1))
        
        if student_evaluations:
            # Öğrenci istatistikleri
            scores = [eval_data['percentage'] for eval_data in student_evaluations]
            
            col1, col2, col3, col4 = st.columns(4)
            
            with col1:
                st.metric("📊 Toplam Değerlendirme", len(student_evaluations))
            
            with col2:
                st.metric("📈 Ortalama Puan", f"{sum(scores)/len(scores):.1f}%")
            
            with col3:
                st.metric("🏆 En Yüksek Puan", f"{max(scores):.1f}%")
            
            with col4:
                st.metric("📉 En Düşük Puan", f"{min(scores):.1f}%")
            
            st.markdown("---")
            
            # Öğrencinin tüm değerlendirmeleri
            st.markdown("#### 📄 Tüm Değerlendirmeler")
            
            for i, eval_data in enumerate(student_evaluations, 1):
                with st.expander(f"{i}. {eval_data['file_name']} - {eval_data['grade']} ({eval_data['percentage']:.1f}%)", expanded=False):
                    
                    col1, col2 = st.columns(2)
                    
                    with col1:
                        st.write(f"**📋 Rubrik:** {eval_data['rubrik_name']}")
                        st.write(f"**📅 Tarih:** {eval_data['created_at'].strftime('%d.%m.%Y %H:%M')}")
                        st.write(f"**💯 Puan:** {eval_data['total_score']}")
                        st.write(f"**📊 Yüzde:** {eval_data['percentage']:.1f}%")
                    
                    with col2:
                        if eval_data.get('assignment_title'):
                            st.write(f"**📝 Ödev:** {eval_data['assignment_title']}")
                        
                        # Kriter başarıları
                        if eval_data.get('evaluation_result', {}).get('criteria_scores'):
                            st.write("**📊 Kriter Puanları:**")
                            for criterion in eval_data['evaluation_result']['criteria_scores']:
                                progress_value = criterion['score'] / criterion['max_score']
                                st.write(f"• {criterion['name']}: {criterion['score']}/{criterion['max_score']}")
                                st.progress(progress_value)

def show_rubric_analysis(db):
    """Rubrik bazlı analizler"""
    st.subheader("📋 Rubrik Analizleri")
    
    # Rubrik listesi
    rubrics = list(db.rubrics.find())
    rubric_names = [r['name'] for r in rubrics]
    
    if not rubric_names:
        st.info("📭 Henüz rubrik bulunamadı.")
        return
    
    # Rubrik seçimi
    selected_rubric = st.selectbox("📋 Rubrik Seçin:", ["Tümü"] + rubric_names)
    
    if selected_rubric == "Tümü":
        # Tüm rubriklerin karşılaştırmalı analizi
        st.markdown("### 📊 Rubrik Karşılaştırma")
        
        rubric_stats_pipeline = [
            {"$group": {
                "_id": "$rubric_name",
                "total_uses": {"$sum": 1},
                "avg_score": {"$avg": "$percentage"},
                "success_rate": {
                    "$avg": {
                        "$cond": [{"$gte": ["$percentage", 60]}, 1, 0]
                    }
                }
            }},
            {"$sort": {"total_uses": -1}}
        ]
        
        rubric_stats = list(db.evaluations.aggregate(rubric_stats_pipeline))
        
        if rubric_stats:
            for rubric_data in rubric_stats:
                col1, col2, col3, col4 = st.columns([2, 1, 1, 1])
                
                with col1:
                    st.write(f"📋 **{rubric_data['_id']}**")
                
                with col2:
                    st.metric("📊 Kullanım", rubric_data['total_uses'])
                
                with col3:
                    st.metric("📈 Ortalama", f"{rubric_data['avg_score']:.1f}%")
                
                with col4:
                    st.metric("🏆 Başarı Oranı", f"{rubric_data['success_rate']*100:.1f}%")
    
    else:
        # Seçili rubrikin detaylı analizi
        st.markdown(f"### 📋 {selected_rubric} - Detaylı Analiz")
        
        # Bu rubrikle yapılan değerlendirmeler
        rubric_evaluations = list(db.evaluations.find({"rubric_name": selected_rubric}))
        
        if rubric_evaluations:
            # Genel istatistikler
            scores = [eval_data['percentage'] for eval_data in rubric_evaluations]
            
            col1, col2, col3, col4 = st.columns(4)
            
            with col1:
                st.metric("📊 Toplam Kullanım", len(rubric_evaluations))
            
            with col2:
                st.metric("📈 Ortalama Puan", f"{sum(scores)/len(scores):.1f}%")
            
            with col3:
                success_count = len([s for s in scores if s >= 60])
                success_rate = (success_count / len(scores)) * 100
                st.metric("🏆 Başarı Oranı", f"{success_rate:.1f}%")
            
            with col4:
                st.metric("📊 En Yüksek", f"{max(scores):.1f}%")
            
            st.markdown("---")
            
            # Kriter bazlı analiz
            st.markdown("#### 📊 Kriter Bazlı Performans")
            
            # İlgili rubriği bul
            rubric_doc = db.rubrics.find_one({"name": selected_rubric})
            
            if rubric_doc:
                # Her kriter için ortalama başarı hesapla
                for criterion in rubric_doc['criteria']:
                    criterion_scores = []
                    
                    for eval_data in rubric_evaluations:
                        if eval_data.get('evaluation_result', {}).get('criteria_scores'):
                            for crit_score in eval_data['evaluation_result']['criteria_scores']:
                                if crit_score['name'] == criterion['name']:
                                    criterion_scores.append(crit_score['score'] / crit_score['max_score'])
                    
                    if criterion_scores:
                        avg_performance = sum(criterion_scores) / len(criterion_scores)
                        
                        col1, col2 = st.columns([3, 1])
                        
                        with col1:
                            st.write(f"**{criterion['name']}** ({criterion['weight']} puan)")
                            st.progress(avg_performance)
                            
                        with col2:
                            st.metric("Ortalama", f"{avg_performance*100:.1f}%")
        else:
            st.info(f"📭 {selected_rubric} rubriği henüz kullanılmamış.")

def show_graphic_analysis(db):
    """Grafik analizler"""
    st.subheader("📈 Grafik Analizler")
    
    st.info("📊 Gelişmiş grafik analizleri için Plotly/Chart.js kütüphaneleri eklenecek.")
    
    # Basit metin tabanlı grafikler
    st.markdown("### 📊 Harf Notu Dağılımı")
    
    grade_pipeline = [
        {"$group": {"_id": "$grade", "count": {"$sum": 1}}},
        {"$sort": {"_id": 1}}
    ]
    grade_distribution = list(db.evaluations.aggregate(grade_pipeline))
    
    if grade_distribution:
        total_count = sum(g['count'] for g in grade_distribution)
        
        for grade_data in grade_distribution:
            grade = grade_data['_id']
            count = grade_data['count']
            percentage = (count / total_count) * 100
            
            # Basit bar chart
            bar_length = int(percentage / 5)  # Her % için bir █
            bar = "█" * bar_length
            
            st.write(f"**{grade}**: {bar} {count} öğrenci ({percentage:.1f}%)")
    
    st.markdown("---")
    
    # Zaman bazlı trend
    st.markdown("### 📅 Zaman Bazlı Trend")
    
    # Son 30 günün değerlendirmeleri
    from datetime import datetime, timedelta
    
    thirty_days_ago = datetime.now() - timedelta(days=30)
    
    recent_evaluations = list(db.evaluations.find({
        "created_at": {"$gte": thirty_days_ago}
    }).sort("created_at", 1))
    
    if recent_evaluations:
        # Günlük ortalama hesapla
        daily_stats = {}
        
        for eval_data in recent_evaluations:
            date_key = eval_data['created_at'].strftime('%Y-%m-%d')
            
            if date_key not in daily_stats:
                daily_stats[date_key] = {'scores': [], 'count': 0}
            
            daily_stats[date_key]['scores'].append(eval_data['percentage'])
            daily_stats[date_key]['count'] += 1
        
        st.write("**Son 30 Günlük Trend:**")
        
        for date_key in sorted(daily_stats.keys()):
            data = daily_stats[date_key]
            avg_score = sum(data['scores']) / len(data['scores'])
            
            # Basit trend gösterimi
            trend_bar = "█" * int(avg_score / 10)
            st.write(f"{date_key}: {trend_bar} {avg_score:.1f}% ({data['count']} değerlendirme)")
    else:
        st.info("📭 Son 30 günde değerlendirme bulunamadı.")

def show_export_options(db):
    """Export seçenekleri"""
    st.subheader("📄 Rapor Export Seçenekleri")
    
    st.markdown("### 📊 Mevcut Veriler")
    
    # Export edilebilir veri türleri
    total_evaluations = db.evaluations.count_documents({})
    total_students = len(db.evaluations.distinct("student_name"))
    total_rubrics = db.rubrics.count_documents({})
    
    col1, col2, col3 = st.columns(3)
    
    with col1:
        st.metric("📄 Değerlendirmeler", total_evaluations)
        if st.button("📥 Değerlendirmeler Export", disabled=True):
            st.info("CSV/Excel export özelliği yakında eklenecek!")
    
    with col2:
        st.metric("👥 Öğrenci Verileri", total_students)
        if st.button("📥 Öğrenci Raporu Export", disabled=True):
            st.info("PDF rapor özelliği yakında eklenecek!")
    
    with col3:
        st.metric("📋 Rubrik Verileri", total_rubrics)
        if st.button("📥 Rubrik Analizi Export", disabled=True):
            st.info("Rubrik analizi export özelliği yakında eklenecek!")
    
    st.markdown("---")
    
    st.markdown("### ⚙️ Export Ayarları")
    
    col1, col2 = st.columns(2)
    
    with col1:
        date_range = st.date_input(
            "📅 Tarih Aralığı",
            value=[datetime.now().date() - timedelta(days=30), datetime.now().date()],
            help="Export edilecek verilerin tarih aralığını seçin"
        )
    
    with col2:
        export_format = st.selectbox(
            "📁 Export Formatı",
            ["PDF", "Excel (.xlsx)", "CSV", "JSON"],
            help="Export edilecek dosya formatını seçin"
        )
    
    include_details = st.checkbox("📋 Detaylı bilgileri dahil et", value=True)
    include_feedback = st.checkbox("💬 AI geri bildirimlerini dahil et", value=False)
    
    if st.button("🚀 Export Oluştur", type="primary", disabled=True):
        st.info(f"🔄 {export_format} formatında export özelliği geliştiriliyor...")
    
    st.markdown("---")
    
    st.markdown("### 📧 Otomatik Raporlama")
    
    st.info("📬 Haftalık/aylık otomatik rapor gönderimi özelliği yakında eklenecek!")
    
    col1, col2 = st.columns(2)
    
    with col1:
        auto_report_frequency = st.selectbox(
            "📅 Rapor Sıklığı",
            ["Haftalık", "Aylık", "Dönemlik"],
            disabled=True
        )
    
    with col2:
        report_recipients = st.text_input(
            "📧 E-posta Adresleri",
            placeholder="ornek@email.com, diger@email.com",
            disabled=True
        )
    
    if st.button("📧 Otomatik Rapor Ayarla", disabled=True):
        st.info("Otomatik e-posta raporu özelliği yakında!")

def grade_converter(percentage):
    """Yüzdeyi harf notuna çevir"""
    if percentage >= 90:
        return "AA"
    elif percentage >= 85:
        return "BA"
    elif percentage >= 75:
        return "BB"
    elif percentage >= 65:
        return "CB"
    elif percentage >= 55:
        return "CC"
    elif percentage >= 45:
        return "DC"
    else:
        return "FF"





if __name__ == "__main__":
    main()
